# -*- coding: utf-8 -*-
"""Markdown to Word converter for brainstorming-mmba-plan.md"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT = os.path.join(SCRIPT_DIR, "..", "brainstorming-mmba-plan.md")
OUTPUT = os.path.join(SCRIPT_DIR, "..", "电销系统对接MMBA工作手机方案.docx")


# ── inline formatting ────────────────────────────────────

def add_inline_runs(paragraph, text, base_bold=False, base_size=Pt(10.5)):
    """Parse **bold**, ~~strike~~, `code` and add runs to paragraph."""
    pattern = r"(\*\*(.+?)\*\*)|(~~(.+?)~~)|(`([^`]+?)`)"
    last = 0
    for m in re.finditer(pattern, text):
        # text before match
        if m.start() > last:
            run = paragraph.add_run(text[last : m.start()])
            run.font.size = base_size
            run.font.bold = base_bold
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if m.group(2):  # bold
            run = paragraph.add_run(m.group(2))
            run.font.bold = True
            run.font.size = base_size
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        elif m.group(4):  # strikethrough
            run = paragraph.add_run(m.group(4))
            run.font.strike = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.size = base_size
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        elif m.group(6):  # code
            run = paragraph.add_run(m.group(6))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        last = m.end()
    # remaining text
    if last < len(text):
        run = paragraph.add_run(text[last:])
        run.font.size = base_size
        run.font.bold = base_bold
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if last == 0 and not text:
        run = paragraph.add_run("")
        run.font.size = base_size


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_font(paragraph, font_name="Microsoft YaHei", size=Pt(10.5)):
    """Set default font for a paragraph."""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = size
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


# ── parse markdown ────────────────────────────────────────

def parse_md(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()
    lines = [l.rstrip("\n") for l in lines]

    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # empty
        if line.strip() == "":
            i += 1
            continue

        # heading
        hm = re.match(r"^(#{1,4})\s+(.*)", line)
        if hm:
            blocks.append(("heading", len(hm.group(1)), hm.group(2)))
            i += 1
            continue

        # code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing
            blocks.append(("code", "\n".join(code_lines)))
            continue

        # table
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split("|")[1:-1]]
                if cells and all(re.match(r"^[-:]+$", c) for c in cells):
                    continue
                if cells:
                    rows.append(cells)
            if rows:
                blocks.append(("table", rows))
            continue

        # blockquote
        if line.strip().startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(re.sub(r"^>\s?", "", lines[i]))
                i += 1
            blocks.append(("quote", "\n".join(quote_lines)))
            continue

        # checkbox list
        cm = re.match(r"^\s*-\s+\[[ x]\]\s+", line)
        if cm:
            items = []
            while i < len(lines) and re.match(r"^\s*-\s+\[[ x]\]\s+", lines[i]):
                checked = "[x]" in lines[i]
                txt = re.sub(r"^\s*-\s+\[[ x]\]\s+", "", lines[i])
                items.append((txt, checked))
                i += 1
            blocks.append(("checklist", items))
            continue

        # unordered list
        lm = re.match(r"^\s*[-*]\s+", line)
        if lm:
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]))
                i += 1
            blocks.append(("list", items))
            continue

        # hr
        if re.match(r"^---+$", line.strip()):
            i += 1
            continue

        # paragraph
        blocks.append(("paragraph", line))
        i += 1

    return blocks


# ── build docx ────────────────────────────────────────────

HEADING_COLORS = {
    1: RGBColor(0x1F, 0x29, 0x37),
    2: RGBColor(0x1F, 0x4E, 0x79),
    3: RGBColor(0x2E, 0x75, 0xB6),
    4: RGBColor(0x40, 0x40, 0x40),
}

HEADING_SIZES = {
    1: Pt(18),
    2: Pt(15),
    3: Pt(13),
    4: Pt(12),
}


def build_docx(blocks, output_path):
    doc = Document()

    # default font
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    for block in blocks:
        btype = block[0]

        if btype == "heading":
            _, level, text = block
            p = doc.add_paragraph()
            add_inline_runs(p, text, base_bold=True, base_size=HEADING_SIZES.get(level, Pt(12)))
            for run in p.runs:
                run.font.color.rgb = HEADING_COLORS.get(level, RGBColor(0, 0, 0))
                run.font.bold = True
            p.paragraph_format.space_before = Pt(16 if level <= 2 else 12)
            p.paragraph_format.space_after = Pt(8)

        elif btype == "paragraph":
            _, text = block
            p = doc.add_paragraph()
            add_inline_runs(p, text)
            p.paragraph_format.space_after = Pt(4)

        elif btype == "quote":
            _, text = block
            for ql in text.split("\n"):
                p = doc.add_paragraph()
                add_inline_runs(p, ql.strip(), base_size=Pt(10))
                for run in p.runs:
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                p.paragraph_format.left_indent = Cm(0.8)
                p.paragraph_format.space_after = Pt(2)
                # left border via XML
                pPr = p._element.get_or_add_pPr()
                borders = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    '  <w:left w:val="single" w:sz="6" w:space="8" w:color="BBBBBB"/>'
                    "</w:pBdr>"
                )
                pPr.append(borders)

        elif btype == "code":
            _, text = block
            for cl in text.split("\n"):
                p = doc.add_paragraph()
                run = p.add_run(cl if cl else " ")
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.left_indent = Cm(0.6)
                # background shading
                pPr = p._element.get_or_add_pPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F5F5F5"/>')
                pPr.append(shading)
            # spacer
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        elif btype == "list":
            _, items = block
            for item in items:
                p = doc.add_paragraph()
                run = p.add_run("•  ")
                run.font.name = "Microsoft YaHei"
                run.font.size = Pt(10.5)
                add_inline_runs(p, item)
                p.paragraph_format.left_indent = Cm(0.8)
                p.paragraph_format.space_after = Pt(2)

        elif btype == "checklist":
            _, items = block
            for item_text, checked in items:
                p = doc.add_paragraph()
                prefix = "☑  " if checked else "☐  "
                run = p.add_run(prefix)
                run.font.name = "Microsoft YaHei"
                run.font.size = Pt(10.5)
                add_inline_runs(p, item_text)
                p.paragraph_format.left_indent = Cm(0.8)
                p.paragraph_format.space_after = Pt(2)

        elif btype == "table":
            _, rows = block
            if not rows:
                continue
            col_count = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=col_count)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"

            for ri, row in enumerate(rows):
                for ci in range(col_count):
                    cell = table.cell(ri, ci)
                    cell_text = row[ci] if ci < len(row) else ""
                    # clear default paragraph
                    cell.paragraphs[0].clear()
                    p = cell.paragraphs[0]
                    is_header = ri == 0
                    add_inline_runs(p, cell_text, base_bold=is_header, base_size=Pt(9.5))
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(2)
                    if is_header:
                        set_cell_shading(cell, "E8EDF2")

            # spacer
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    doc.save(output_path)
    print(f"Done! Saved to: {output_path}")


# ── main ──────────────────────────────────────────────────

if __name__ == "__main__":
    print(f"Reading: {INPUT}")
    blocks = parse_md(INPUT)
    print(f"Parsed {len(blocks)} blocks")
    print("Generating Word document...")
    build_docx(blocks, OUTPUT)
